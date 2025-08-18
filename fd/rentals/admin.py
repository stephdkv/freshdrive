from django.contrib import admin
from django.conf import settings
from .models import Transport, RentalApplication, Client, Calendar, Advantages, TransportImage, Blog
from .forms import RentalApplicationForm
from django.http import HttpResponse, JsonResponse
from docx import Document
from django.template.defaultfilters import date as _date
from django.contrib.auth.models import Group, Permission, User
from django.contrib.auth.admin import UserAdmin
from django.core.exceptions import PermissionDenied
from django.db.models import Sum, Q, Avg
import io
import os
from datetime import datetime, date, timedelta
from django.utils.html import format_html
from django.urls import path
from django.shortcuts import render
from urllib.parse import quote
from django.contrib.admin.views.decorators import staff_member_required
import json
from django.db import models
from django_summernote.admin import SummernoteModelAdmin
from .models import UserProfile
from .models import Review

# Отменяем регистрацию стандартного UserAdmin
admin.site.unregister(User)

class UserProfileInline(admin.StackedInline):
    model = UserProfile
    can_delete = False
    verbose_name_plural = 'Профиль пользователя'

class CustomUserAdmin(UserAdmin):
    inlines = (UserProfileInline,)
    list_display = ('username', 'email', 'first_name', 'last_name', 'is_staff', 'get_groups')
    list_filter = ('is_staff', 'is_superuser', 'groups')
    
    def get_groups(self, obj):
        return ", ".join([group.name for group in obj.groups.all()])
    get_groups.short_description = "Группы"

    def get_fieldsets(self, request, obj=None):
        if not obj:
            return self.add_fieldsets
            
        if request.user.is_superuser:
            return (
                (None, {'fields': ('username', 'password')}),
                ('Персональная информация', {'fields': ('first_name', 'last_name', 'email')}),
                ('Права доступа', {
                    'fields': ('is_active', 'is_staff', 'is_superuser', 'groups'),
                }),
            )
        return (
            (None, {'fields': ('username', 'password')}),
            ('Персональная информация', {'fields': ('first_name', 'last_name', 'email')}),
            ('Права доступа', {
                'fields': ('is_active', 'groups'),
            }),
        )

    def get_readonly_fields(self, request, obj=None):
        base = super().get_readonly_fields(request, obj)
        extra = ('total_cost_display', 'rental_days_display', 'daily_rate_display')
        if obj:
            if isinstance(base, (list, tuple)):
                return tuple(base) + extra
            return extra
        return base

    def save_formset(self, request, form, formset, change):
        """Сохраняем инлайн-формы"""
        instances = formset.save(commit=False)
        for instance in instances:
            instance.save()
        formset.save_m2m()
        super().save_formset(request, form, formset, change)

# Удаляем стандартную регистрацию User, если есть
try:
    admin.site.unregister(User)
except admin.sites.NotRegistered:
    pass
# Регистрируем User с новым админом
admin.site.register(User, CustomUserAdmin)

# Отменяем регистрацию стандартного GroupAdmin, так как он нам не нужен
# admin.site.unregister(Group)

template_path = os.path.join(settings.BASE_DIR, 'rentals', 'templates', 'contracts', 'rental_contract_template.docx')
doc = Document(template_path)

def create_manager_group():
    """Создает группу менеджеров с необходимыми правами если она не существует"""
    group, created = Group.objects.get_or_create(name='Manager')
    if created:
        # Права для Transport
        transport_permissions = Permission.objects.filter(
            content_type__app_label='rentals',
            content_type__model='transport',
            codename__in=['view_transport']
        )
        # Права для RentalApplication
        rental_permissions = Permission.objects.filter(
            content_type__app_label='rentals',
            content_type__model='rentalapplication',
            codename__in=['add_rentalapplication', 'change_rentalapplication', 'view_rentalapplication']
        )
        
        # Добавляем все разрешения в группу
        group.permissions.set(list(transport_permissions) + list(rental_permissions))

class OverdueStatusFilter(admin.SimpleListFilter):
    title = 'Просрочена'
    parameter_name = 'overdue'

    def lookups(self, request, model_admin):
        return (
            ('yes', 'Только просроченные'),
            ('no', 'Без просроченных'),
        )

    def queryset(self, request, queryset):
        if self.value() == 'yes':
            return queryset.filter(status='active').filter(rental_end_date__lt=datetime.now().date())
        if self.value() == 'no':
            return queryset.exclude(status='active', rental_end_date__lt=datetime.now().date())
        return queryset

class TransportImageInline(admin.TabularInline):
    model = TransportImage
    extra = 1

@admin.register(Transport)
class TransportAdmin(SummernoteModelAdmin):
    summernote_fields = ('description',)
    inlines = [TransportImageInline]
    list_display = (
        'number', 'name', 'model', 'year', 'color', 'registration_number', 'vin_number',
        'price_per_day', 'price_3_6_days', 'price_7_30_days', 'price_30_plus_days', 'city', 'category', 'transmission'
    )
    search_fields = ('number', 'name', 'model', 'registration_number', 'color', 'vin_number', 'slug')
    list_filter = ('year', 'city', 'category', 'transmission', 'condition')
    prepopulated_fields = {'slug': ('name',)}
    ordering = ('number',)

    fieldsets = (
        ('Основная информация', {
            'fields': ('number', 'name', 'slug', 'model', 'year', 'color', 'registration_number', 'vin_number', 'city', 'category')
        }),
        ('Технические характеристики', {
            'fields': ('power_hp', 'transmission', 'condition', 'mileage_km', 'engine_volume_l')
        }),
        ('Изображения', {
            'fields': ('main_image', 'thumbnail_image')
        }),
        ('Описание', {
            'fields': ('description',)
        }),
        ('Цены', {
            'fields': ('price_per_day', 'price_3_6_days', 'price_7_30_days', 'price_30_plus_days')
        }),
    )

    def has_add_permission(self, request):
        return request.user.is_superuser

    def has_change_permission(self, request, obj=None):
        return request.user.is_superuser

    def has_delete_permission(self, request, obj=None):
        return request.user.is_superuser

    def save_model(self, request, obj, form, change):
        if not request.user.is_superuser:
            profile = getattr(request.user, 'profile', None)
            if profile:
                obj.city = profile.city
        super().save_model(request, obj, form, change)

@admin.register(RentalApplication)
class RentalApplicationAdmin(admin.ModelAdmin):
    form = RentalApplicationForm
    list_display = ('get_colored_status', 'full_name', 'phone_number', 'transport', 'rental_start_date', 
                   'rental_end_date', 'rental_days_display', 'get_rate_type_display',
                   'daily_rate_display', 'get_discount_display', 'get_discount_amount_display',
                   'get_security_deposit_display', 'get_total_cost_display', 'how_did_you_find_us',
                   'manager_display',
                   )
    list_filter = ('status', 'rental_start_date', 'rental_end_date', 'created_at', 'discount', 'how_did_you_find_us', OverdueStatusFilter)
    search_fields = ('full_name', 'phone_number', 'passport_number', 
                    'transport__name', 'transport__model')
    date_hierarchy = 'rental_start_date'
    list_per_page = 30
    
    def get_fieldsets(self, request, obj=None):
        base_fields = (
            'rental_start_date', 'rental_end_date', 'transport', 'security_deposit', 'discount', 'status',
        )
        if obj:
            details_fields = base_fields + (
                'total_cost_display', 'rental_days_display', 'daily_rate_display',
            )
        else:
            details_fields = base_fields
        return (
            ('Информация об арендаторе', {
                'fields': ('client', 'full_name', 'phone_number', 'how_did_you_find_us')
            }),
            ('Детали аренды', {
                'fields': details_fields
            }),
            ('Паспортные данные', {
                'fields': ('passport_number', 'passport_issued_by', 'passport_issue_date')
            }),
        )

    readonly_fields = ('total_cost_display', 'rental_days_display', 'daily_rate_display')

    def total_cost_display(self, obj):
        return format_html('<span style="color: #28a745; font-weight: bold;">{} ₽</span>', f"{obj.calculate_total_cost():,}")
    total_cost_display.short_description = 'Общая стоимость'

    def rental_days_display(self, obj):
        return format_html('<span style="color: #28a745; font-weight: bold;">{} суток</span>', obj.get_rental_days())
    rental_days_display.short_description = 'Количество суток'

    def daily_rate_display(self, obj):
        return format_html('<span style="color: #28a745; font-weight: bold;">{} ₽/день</span>', f"{obj.get_daily_rate():,}")
    daily_rate_display.short_description = 'Стоимость в сутки'

    def get_rate_type_display(self, obj):
        return obj.get_rate_type()
    get_rate_type_display.short_description = 'Тип тарифа'

    def get_discount_display(self, obj):
        return f"{obj.discount}%"
    get_discount_display.short_description = 'Скидка'

    def get_discount_amount_display(self, obj):
        return f"{obj.get_discount_amount():,} ₽"
    get_discount_amount_display.short_description = 'Сумма скидки'

    def get_total_cost_display(self, obj):
        return f"{obj.calculate_total_cost():,} ₽"
    get_total_cost_display.short_description = 'Итого со скидкой'

    def get_security_deposit_display(self, obj):
        return f"{int(obj.security_deposit):,} ₽"
    get_security_deposit_display.short_description = 'Залог'
    
    def make_active(self, request, queryset):
        from django.utils import timezone
        for application in queryset:
            if request.user.is_superuser or application.can_change_status_to(application.STATUS_ACTIVE):
                application.status = application.STATUS_ACTIVE
                application.activated_at = timezone.now()
                application.save()
                self.message_user(request, f'Заявка "{application}" переведена в статус "Активная"', level='success')
            else:
                old_status_display = dict(application.STATUS_CHOICES).get(application.status, application.status)
                new_status_display = dict(application.STATUS_CHOICES).get(application.STATUS_ACTIVE, application.STATUS_ACTIVE)
                self.message_user(request, f'Нельзя сменить статус заявки "{application}" с "{old_status_display}" на "{new_status_display}"', level='error')
    make_active.short_description = "Сделать активными"

    def make_completed(self, request, queryset):
        for application in queryset:
            if request.user.is_superuser or application.can_change_status_to(application.STATUS_COMPLETED):
                application.status = application.STATUS_COMPLETED
                application.save()
                self.message_user(request, f'Заявка "{application}" переведена в статус "Завершенная"', level='success')
            else:
                old_status_display = dict(application.STATUS_CHOICES).get(application.status, application.status)
                new_status_display = dict(application.STATUS_CHOICES).get(application.STATUS_COMPLETED, application.STATUS_COMPLETED)
                self.message_user(request, f'Нельзя сменить статус заявки "{application}" с "{old_status_display}" на "{new_status_display}"', level='error')
    make_completed.short_description = "Отметить как завершенные"

    def make_cancelled(self, request, queryset):
        for application in queryset:
            if request.user.is_superuser or application.can_change_status_to(application.STATUS_CANCELLED):
                application.status = application.STATUS_CANCELLED
                application.save()
                self.message_user(request, f'Заявка "{application}" переведена в статус "Отмененная"', level='success')
            else:
                old_status_display = dict(application.STATUS_CHOICES).get(application.status, application.status)
                new_status_display = dict(application.STATUS_CHOICES).get(application.STATUS_CANCELLED, application.STATUS_CANCELLED)
                self.message_user(request, f'Нельзя сменить статус заявки "{application}" с "{old_status_display}" на "{new_status_display}"', level='error')
    make_cancelled.short_description = "Отметить как отмененные"

    def complete_early_and_print_addendum(self, request, queryset):
        """Досрочно завершить аренду и распечатать Дополнение к договору (return.docx)"""
        if queryset.count() != 1:
            self.message_user(request, "Выберите одну активную заявку для досрочного завершения.", level='error')
            return
        application = queryset.first()
        if application.status != RentalApplication.STATUS_ACTIVE:
            self.message_user(request, "Досрочно завершить можно только активную аренду!", level='error')
            return
        try:
            application.complete_early()
        except Exception as e:
            self.message_user(request, f"Ошибка: {e}", level='error')
            return
        return self.generate_return_addendum(request, application)
    complete_early_and_print_addendum.short_description = "Досрочно завершить и распечатать Дополнение"

    def generate_return_addendum(self, request, application):
        """Генерирует и возвращает Дополнение к договору (return.docx)"""
        import io
        from docx import Document
        import os
        from django.conf import settings
        from django.http import HttpResponse
        from django.template.defaultfilters import date as _date
        from datetime import datetime
        # Путь к шаблону
        template_path = os.path.join(settings.BASE_DIR, 'rentals', 'templates', 'contracts', 'return.docx')
        doc = Document(template_path)
        transport = application.transport
        # Фактические значения для доп. соглашения
        actual_return_date = application.rental_end_date
        actual_rental_days = application.get_rental_days()
        actual_cost = application.calculate_total_cost()
        # Для возврата: если был внесен депозит, и actual_cost < изначальной суммы, считаем возврат
        # (допустим, изначальная сумма = application.total_cost_before_early, если нужно — добавить в модель)
        # Здесь просто считаем разницу между старой и новой суммой, если нужно
        refund_amount = 0
        if application.original_total_cost:
            refund_amount = max(0, int(application.original_total_cost) - int(actual_cost))
        # Если нет — не выводим
        context = {
            'full_name': application.full_name,
            'phone_number': application.phone_number,
            'color': transport.color,
            'passport_number': application.passport_number,
            'passport_issued_by': application.passport_issued_by,
            'passport_issue_date': _date(application.passport_issue_date, "d.m.Y"),
            'rental_start_date': _date(application.rental_start_date, "d.m.Y"),
            'rental_end_date': _date(application.rental_end_date, "d.m.Y"),
            'transport_model': f"№{transport.number} - {transport.name} {transport.model}",
            'registration_number': transport.registration_number,
            'vin_number': transport.vin_number,
            'rental_days': application.get_rental_days(),
            'rate_type': application.get_rate_type(),
            'daily_rate': application.get_daily_rate(),
            'total_cost': application.calculate_total_cost(),
            'security_deposit': int(application.security_deposit),
            'today_date': datetime.now().strftime("%d.%m.%Y"),
            # Новые поля:
            'actual_return_date': _date(actual_return_date, "d.m.Y"),
            'actual_rental_days': actual_rental_days,
            'actual_cost': actual_cost,
            'refund_amount': refund_amount,
        }
        # Замена плейсхолдеров
        for paragraph in doc.paragraphs:
            text = paragraph.text
            for key, value in context.items():
                placeholder = f"{{{{ {key} }}}}"
                if placeholder in text:
                    text = text.replace(placeholder, str(value))
            if text != paragraph.text:
                for run in paragraph.runs:
                    run.text = ""
                if paragraph.runs:
                    paragraph.runs[0].text = text
                else:
                    paragraph.add_run(text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text
                        for key, value in context.items():
                            placeholder = f"{{{{ {key} }}}}"
                            if placeholder in text:
                                text = text.replace(placeholder, str(value))
                        if text != paragraph.text:
                            for run in paragraph.runs:
                                run.text = ""
                            if paragraph.runs:
                                paragraph.runs[0].text = text
                            else:
                                paragraph.add_run(text)
        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        filename = f"Дополнение к договору - {application.full_name} от {datetime.now().strftime('%d.%m.%Y')}.docx"
        encoded_filename = filename.encode('utf-8').decode('latin-1')
        response = HttpResponse(f.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename=\"{encoded_filename}\"'
        return response

    def generate_contract(self, request, queryset):
        # Проверяем права на генерацию договора
        if not (request.user.is_superuser or request.user.groups.filter(name='Manager').exists()):
            raise PermissionDenied("У вас нет прав на генерацию договора")

        if queryset.count() != 1:
            self.message_user(request, "Выберите одну заявку для генерации договора.", level='error')
            return

        application = queryset.first()
        transport = application.transport

        # Загрузка шаблона
        template_path = os.path.join(settings.BASE_DIR, 'rentals', 'templates', 'contracts', 'rental_contract_template.docx')
        doc = Document(template_path)

        # Замена всех {{ ... }} на реальные значения
        context = {
            'full_name': application.full_name,
            'phone_number': application.phone_number,
            'color': transport.color,
            'passport_number': application.passport_number,
            'passport_issued_by': application.passport_issued_by,
            'passport_issue_date': _date(application.passport_issue_date, "d.m.Y"),
            'rental_start_date': _date(application.rental_start_date, "d.m.Y"),
            'rental_end_date': _date(application.rental_end_date, "d.m.Y"),
            'transport_model': f"№{transport.number} - {transport.name} {transport.model}",
            'registration_number': transport.registration_number,
            'vin_number': transport.vin_number,
            'rental_days': application.get_rental_days(),
            'rate_type': application.get_rate_type(),
            'daily_rate': application.get_daily_rate(),
            'total_cost': application.calculate_total_cost(),
            'security_deposit': int(application.security_deposit),
            'today_date': datetime.now().strftime("%d.%m.%Y"),
        }

        # Проходим по всем параграфам
        for paragraph in doc.paragraphs:
            text = paragraph.text
            for key, value in context.items():
                placeholder = f"{{{{ {key} }}}}"
                if placeholder in text:
                    text = text.replace(placeholder, str(value))
            if text != paragraph.text:
                for run in paragraph.runs:
                    run.text = ""
                if paragraph.runs:
                    paragraph.runs[0].text = text
                else:
                    paragraph.add_run(text)

        # Обработка таблиц
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text
                        for key, value in context.items():
                            placeholder = f"{{{{ {key} }}}}"
                            if placeholder in text:
                                text = text.replace(placeholder, str(value))
                        if text != paragraph.text:
                            for run in paragraph.runs:
                                run.text = ""
                            if paragraph.runs:
                                paragraph.runs[0].text = text
                            else:
                                paragraph.add_run(text)

        # Сохранение во временный файл и отдача пользователю
        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        filename = f"Договор аренды - {application.full_name} от {datetime.now().strftime('%d.%m.%Y')}.docx"
        encoded_filename = filename.encode('utf-8').decode('latin-1')
        response = HttpResponse(f.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="{encoded_filename}"'
        return response
    generate_contract.short_description = "Печать договора аренды"

    actions = ['generate_contract', 'make_active', 'make_completed', 'make_cancelled', 'complete_early_and_print_addendum']

    def get_colored_status(self, obj):
        status_classes = {
            'reserved': 'status-reserved',
            'active': 'status-active',
            'completed': 'status-completed',
            'cancelled': 'status-cancelled',
            'overdue': 'status-overdue',
        }
        # Виртуальный статус 'Просрочена'
        if obj.is_overdue:
            return format_html(
                '<span class="status-badge {}" style="background:#dc3545; color:white;">{}</span>',
                status_classes['overdue'],
                'Просрочена'
            )
        return format_html(
            '<span class="status-badge {}">{}</span>',
            status_classes.get(obj.status, ''),
            obj.get_status_display()
        )
    get_colored_status.short_description = 'Статус'
    get_colored_status.admin_order_field = 'status'

    list_filter = ('status', 'rental_start_date', 'rental_end_date', 'created_at', 'discount', 'how_did_you_find_us', OverdueStatusFilter)

    def changelist_view(self, request, extra_context=None):
        # Get the filtered queryset
        response = super().changelist_view(request, extra_context=extra_context)
        
        try:
            # Get the filtered queryset
            qs = response.context_data['cl'].queryset
            
            # Calculate totals
            total_cost = sum(app.calculate_total_cost() for app in qs)
            total_security_deposit = sum(int(app.security_deposit or 0) for app in qs)
            total_days = sum(app.get_rental_days() for app in qs)
            
            # Add totals to the context
            response.context_data['total_cost'] = f"{total_cost:,} ₽"
            response.context_data['total_security_deposit'] = f"{total_security_deposit:,} ₽"
            response.context_data['total_days'] = f"{total_days} дн."
            
        except (AttributeError, KeyError):
            # If there's an error, just return the response without totals
            return response
            
        return response

    def get_urls(self):
        from django.urls import path
        urls = super().get_urls()
        custom_urls = [
            path('analytics/', self.admin_site.admin_view(self.analytics_view), name='rentalapplication_analytics'),
        ]
        return custom_urls + urls

    def analytics_view(self, request):
        from django.db.models import Count, Sum, Avg, Q
        from .models import RentalApplication, Transport, Client
        # Пример простой аналитики
        total = RentalApplication.objects.count()
        by_status_raw = list(RentalApplication.objects.values('status').annotate(count=Count('id')).order_by('status'))
        status_map = dict(RentalApplication.STATUS_CHOICES)
        by_status = [
            {'status': status_map.get(item['status'], item['status']), 'count': item['count']}
            for item in by_status_raw
        ]
        total_sum = RentalApplication.objects.aggregate(total=Sum('security_deposit'))['total']
        # Топ-10 клиентов по количеству заявок
        top_clients = Client.objects.annotate(num=Count('rental_applications')).order_by('-num')[:10]
        # Динамика по месяцам
        months = []
        now = date.today()
        for i in range(11, -1, -1):
            month = (now.replace(day=1) - timedelta(days=30*i)).replace(day=1)
            months.append(month)
        month_labels = [m.strftime('%Y-%m') for m in months]
        month_counts = [RentalApplication.objects.filter(created_at__year=m.year, created_at__month=m.month).count() for m in months]
        # 1. Средняя продолжительность аренды (в днях)
        all_apps = RentalApplication.objects.exclude(rental_start_date__isnull=True).exclude(rental_end_date__isnull=True)
        if all_apps.exists():
            days_list = [(app.rental_end_date - app.rental_start_date).days for app in all_apps if app.rental_end_date and app.rental_start_date]
            mean_days = round(sum(days_list) / len(days_list), 1) if days_list else 0
        else:
            mean_days = 0
        # 2. Средний чек
        all_apps_with_cost = RentalApplication.objects.exclude(rental_start_date__isnull=True).exclude(rental_end_date__isnull=True)
        if all_apps_with_cost.exists():
            costs_list = []
            for app in all_apps_with_cost:
                cost = app.calculate_total_cost()
                if cost is not None and cost > 0:
                    costs_list.append(cost)
            mean_total_cost = round(sum(costs_list) / len(costs_list), 0) if costs_list else 0
        else:
            mean_total_cost = 0
        # 3. Топ-5 транспортов
        top_transports = list(RentalApplication.objects.values('transport__name', 'transport__model').annotate(num=Count('id')).order_by('-num')[:5])
        # 4. Источники заявок
        sources = list(RentalApplication.objects.values('how_did_you_find_us').annotate(num=Count('id')).order_by('-num'))
        sources_map = dict(getattr(RentalApplication, 'HOW_DID_YOU_FIND_US_CHOICES', []))
        for s in sources:
            code = s.get('how_did_you_find_us')
            if code == 'friends':
                s['how_did_you_find_us'] = 'От друзей / знакомых'
            elif code == 'internet':
                s['how_did_you_find_us'] = 'Из интернета / через поиск (Google, Яндекс)'
            elif code == 'ads':
                s['how_did_you_find_us'] = 'Увидел рекламу'
            elif code == 'repeat_customer':
                s['how_did_you_find_us'] = 'У вас брал услугу раньше / уже был клиентом'
            elif code == 'catalog':
                s['how_did_you_find_us'] = 'Нашли вас в каталоге / на картах (Google Maps, 2ГИС, Яндекс.Справочник)'
            elif code == 'other':
                s['how_did_you_find_us'] = 'Другое'
            elif code in (None, '', 'none', 'null'):
                s['how_did_you_find_us'] = 'Не указано'
            else:
                s['how_did_you_find_us'] = str(code)
        # 5. Динамика новых клиентов
        new_clients_counts = []
        for m in months:
            count = Client.objects.filter(created_at__year=m.year, created_at__month=m.month).count()
            new_clients_counts.append(count)
        new_clients_months = month_labels
        # 6. Процент заявок со скидкой
        discount_count = RentalApplication.objects.exclude(discount=0).count()
        discount_percent = round((discount_count / total) * 100, 1) if total else 0
        no_discount_count = total - discount_count
        now = datetime.now()
        month_apps = RentalApplication.objects.filter(
            created_at__year=now.year, created_at__month=now.month
        )
        month_earned = month_apps.aggregate(total=models.Sum('original_total_cost'))['total'] or 0
        month_count = month_apps.count()
        context = dict(
            self.admin_site.each_context(request),
            total=total,
            by_status=json.dumps(by_status, ensure_ascii=False),
            total_sum=total_sum,
            top_clients=top_clients,
            month_labels=json.dumps(month_labels, ensure_ascii=False),
            month_counts=json.dumps(month_counts, ensure_ascii=False),
            mean_days=mean_days,
            mean_total_cost=mean_total_cost,
            top_transports=json.dumps(top_transports, ensure_ascii=False),
            sources=json.dumps(sources, ensure_ascii=False),
            new_clients_months=json.dumps(new_clients_months, ensure_ascii=False),
            new_clients_counts=json.dumps(new_clients_counts, ensure_ascii=False),
            discount_percent=discount_percent,
            discount_count=discount_count,
            total_count=total,
            no_discount_count=no_discount_count,
            month_earned=month_earned,
            month_count=month_count,
            opts=self.model._meta,
            title='Аналитика по заявкам',
        )
        from django.shortcuts import render
        return render(request, 'admin/rentals/rentalapplication/analytics.html', context)

    class Media:
        css = {
            'all': (
                'admin/css/forms.css',
                'admin/css/widgets.css',
                'rentals/css/admin.css',
            )
        }
        js = (
            'admin/js/jquery.init.js',
            'admin/js/core.js',
            'admin/js/calendar.js',
            'admin/js/admin/DateTimeShortcuts.js',
            'rentals/js/rental_form.js',
        )

    def get_queryset(self, request):
        qs = super().get_queryset(request)
        if request.user.is_superuser:
            return qs
        profile = getattr(request.user, 'profile', None)
        if profile:
            return qs.filter(city=profile.city)
        return qs.none()

    def save_model(self, request, obj, form, change):
        from django.utils import timezone
        obj._status_error = False
        if change and 'status' in form.changed_data:
            old_obj = type(obj).objects.get(pk=obj.pk)
            old_status = old_obj.status
            new_status = form.cleaned_data['status']
            # Получаем русские названия статусов
            old_status_display = dict(obj.STATUS_CHOICES).get(old_status, old_status)
            new_status_display = dict(obj.STATUS_CHOICES).get(new_status, new_status)
            if not request.user.is_superuser:
                if not old_obj.can_change_status_to(new_status):
                    self.message_user(request, f'Нельзя сменить статус с "{old_status_display}" на "{new_status_display}"', level='error')
                    obj._status_error = True
                    return
            # Если переводим в Активная — сохраняем время активации
            if new_status == obj.STATUS_ACTIVE:
                obj.activated_at = timezone.now()
        if not change:
            obj.created_by = request.user
            profile = getattr(request.user, 'profile', None)
            if profile:
                obj.city = profile.city
            if obj.status == obj.STATUS_ACTIVE:
                obj.activated_at = timezone.now()
        elif not obj.created_by:
            obj.created_by = request.user
        super().save_model(request, obj, form, change)

    def response_change(self, request, obj):
        if hasattr(obj, '_status_error') and obj._status_error:
            from django.http import HttpResponseRedirect
            return HttpResponseRedirect(request.path)
        return super().response_change(request, obj)

    def formfield_for_foreignkey(self, db_field, request, **kwargs):
        if db_field.name == "transport" and not request.user.is_superuser:
            profile = getattr(request.user, 'profile', None)
            from .models import Transport
            if profile:
                kwargs["queryset"] = Transport.objects.filter(city=profile.city)
            else:
                kwargs["queryset"] = Transport.objects.none()
        return super().formfield_for_foreignkey(db_field, request, **kwargs)

    def manager_display(self, obj):
        if obj.created_by:
            if obj.created_by.first_name and obj.created_by.last_name:
                return f"{obj.created_by.first_name} {obj.created_by.last_name}"
            else:
                return obj.created_by.username
        return '-'
    manager_display.short_description = 'Менеджер'

    def get_form(self, request, obj=None, **kwargs):
        Form = super().get_form(request, obj, **kwargs)
        class ModelFormWithRequest(Form):
            def __new__(cls, *args, **kw):
                kw['request'] = request
                return Form(*args, **kw)
        return ModelFormWithRequest

@admin.register(Client)
class ClientAdmin(admin.ModelAdmin):
    list_display = ('full_name', 'phone_number', 'get_passport_info', 'get_rental_count', 'get_last_rental_date', 'created_at')
    search_fields = ('full_name', 'phone_number', 'passport_number')
    ordering = ('phone_number',)
    
    def get_passport_info(self, obj):
        if obj.passport_number:
            return f"Паспорт: {obj.passport_number}"
        return "Паспорт не указан"
    get_passport_info.short_description = 'Паспортные данные'
    
    def get_rental_count(self, obj):
        return obj.rental_applications.count()
    get_rental_count.short_description = 'Количество заявок'
    
    def get_last_rental_date(self, obj):
        last_rental = obj.rental_applications.order_by('-rental_start_date').first()
        if last_rental:
            return last_rental.rental_start_date
        return '-'
    get_last_rental_date.short_description = 'Последняя аренда'

    def create_rental_for_client(self, request, queryset):
        from django.shortcuts import redirect
        from urllib.parse import quote
        
        if queryset.count() == 1:
            # Если выбран только один клиент, перенаправляем на форму создания
            client = queryset.first()
            
            # Формируем URL с предзаполненными данными клиента, включая паспортные данные
            params = {
                'client_id': client.id,
                'full_name': quote(client.full_name),
                'phone_number': quote(str(client.phone_number)),
            }
            
            # Добавляем паспортные данные, если они есть
            if client.passport_number:
                params['passport_number'] = quote(client.passport_number)
            if client.passport_issued_by:
                params['passport_issued_by'] = quote(client.passport_issued_by)
            if client.passport_issue_date:
                params['passport_issue_date'] = client.passport_issue_date.strftime('%Y-%m-%d')
            # Добавляем how_did_you_find_us, если есть
            if client.how_did_you_find_us:
                params['how_did_you_find_us'] = quote(client.how_did_you_find_us)
            
            # Собираем URL с параметрами
            param_string = '&'.join([f'{k}={v}' for k, v in params.items()])
            add_url = f'/admin/rentals/rentalapplication/add/?{param_string}'
            
            return redirect(add_url)
        else:
            # Если выбрано несколько клиентов, показываем сообщение
            self.message_user(request, 'Для создания заявки выберите только одного клиента', level='WARNING')
            return None
    
    create_rental_for_client.short_description = "Создать заявку для выбранного клиента"

    actions = ['create_rental_for_client']

    fieldsets = (
        ('Основная информация', {
            'fields': ('full_name', 'phone_number', 'how_did_you_find_us')
        }),
        ('Паспортные данные', {
            'fields': ('passport_number', 'passport_issued_by', 'passport_issue_date'),
            'classes': ('collapse',)
        }),
        ('История заявок', {
            'fields': ('get_rental_history',),
            'classes': ('collapse',)
        }),
    )
    readonly_fields = ('get_rental_history',)

    def get_urls(self):
        from django.urls import path
        urls = super().get_urls()
        custom_urls = [
            path(
                '<int:client_id>/create-rental/',
                self.admin_site.admin_view(self.create_rental_view),
                name='rentals_client_create_rental',
            ),
            path('analytics/', self.admin_site.admin_view(self.analytics_view), name='client_analytics'),
        ]
        return custom_urls + urls

    def create_rental_view(self, request, client_id):
        from django.shortcuts import redirect
        from django.contrib import messages
        
        try:
            client = Client.objects.get(id=client_id)
            
            # Формируем URL с предзаполненными данными клиента, включая паспортные данные
            params = {
                'client_id': client.id,
                'full_name': quote(client.full_name),
                'phone_number': quote(str(client.phone_number)),
            }
            
            # Добавляем паспортные данные, если они есть
            if client.passport_number:
                params['passport_number'] = quote(client.passport_number)
            if client.passport_issued_by:
                params['passport_issued_by'] = quote(client.passport_issued_by)
            if client.passport_issue_date:
                params['passport_issue_date'] = client.passport_issue_date.strftime('%Y-%m-%d')
            # Добавляем how_did_you_find_us, если есть
            if client.how_did_you_find_us:
                params['how_did_you_find_us'] = quote(client.how_did_you_find_us)
            
            # Собираем URL с параметрами
            param_string = '&'.join([f'{k}={v}' for k, v in params.items()])
            add_url = f'/admin/rentals/rentalapplication/add/?{param_string}'
            
            return redirect(add_url)
            
        except Client.DoesNotExist:
            messages.error(request, 'Клиент не найден')
            return redirect('/admin/rentals/client/')
        except Exception as e:
            messages.error(request, f'Ошибка при создании заявки: {str(e)}')
            return redirect('/admin/rentals/client/')

    def get_rental_history(self, obj):
        rentals = obj.rental_applications.all().order_by('-rental_start_date')
        if not rentals:
            return "Нет заявок"
        
        html = '<table style="width: 100%; border-collapse: collapse;">'
        html += '''
            <tr style="background-color: #f8f9fa;">
                <th style="padding: 8px; border: 1px solid #ddd;">Статус</th>
                <th style="padding: 8px; border: 1px solid #ddd;">Транспорт</th>
                <th style="padding: 8px; border: 1px solid #ddd;">Период</th>
                <th style="padding: 8px; border: 1px solid #ddd;">Сумма</th>
                <th style="padding: 8px; border: 1px solid #ddd;">Действия</th>
            </tr>
        '''
        
        for rental in rentals:
            status_class = {
                'reserved': 'status-reserved',
                'active': 'status-active',
                'completed': 'status-completed',
                'cancelled': 'status-cancelled',
            }.get(rental.status, '')
            
            html += f'''
                <tr>
                    <td style="padding: 8px; border: 1px solid #ddd;">
                        <span class="status-badge {status_class}">{rental.get_status_display()}</span>
                    </td>
                    <td style="padding: 8px; border: 1px solid #ddd;">{rental.transport}</td>
                    <td style="padding: 8px; border: 1px solid #ddd;">
                        {rental.rental_start_date.strftime('%d.%m.%Y')} - {rental.rental_end_date.strftime('%d.%m.%Y')}
                    </td>
                    <td style="padding: 8px; border: 1px solid #ddd;">{rental.calculate_total_cost():,} ₽</td>
                    <td style="padding: 8px; border: 1px solid #ddd;">
                        <a href="/admin/rentals/rentalapplication/{rental.id}/change/" class="button" style=" background: #417690; color: white; text-decoration: none; border-radius: 4px;">Просмотр</a>
                    </td>
                </tr>
            '''
        
        html += '</table>'
        
        # Добавляем кнопку создания новой заявки
        create_button = f'''
            <div style="margin-top: 15px; padding: 10px; background-color: #f8f9fa; border: 1px solid #ddd; border-radius: 4px;">
                <a href="/admin/rentals/client/{obj.id}/create-rental/" 
                   class="button" 
                   style=" background: #28a745; color: white; text-decoration: none; border-radius: 4px; font-weight: bold;">
                    ➕ Создать новую заявку
                </a>
            </div>
        '''
        
        return format_html(html + create_button)
    get_rental_history.short_description = 'История заявок'

    def analytics_view(self, request):
        from django.db.models import Count
        from .models import Client
        import datetime
        import json
        # Всего клиентов
        total = Client.objects.count()
        # С паспортом
        with_passport = Client.objects.exclude(passport_number__isnull=True).exclude(passport_number='').count()
        # Топ-10 по количеству заявок
        top_clients = Client.objects.annotate(num=Count('rental_applications')).order_by('-num')[:10]
        # Динамика по месяцам (количество новых клиентов)
        months = []
        now = date.today()
        for i in range(11, -1, -1):
            month = (now.replace(day=1) - timedelta(days=30*i)).replace(day=1)
            months.append(month)
        month_labels = [m.strftime('%Y-%m') for m in months]
        month_counts = [Client.objects.filter(created_at__year=m.year, created_at__month=m.month).count() for m in months]
        context = dict(
            self.admin_site.each_context(request),
            total=total,
            with_passport=with_passport,
            top_clients=top_clients,
            month_labels=json.dumps(month_labels, ensure_ascii=False),
            month_counts=json.dumps(month_counts, ensure_ascii=False),
            opts=self.model._meta,
            title='Аналитика по клиентам',
        )
        from django.shortcuts import render
        return render(request, 'admin/rentals/client/analytics.html', context)

    def get_queryset(self, request):
        return super().get_queryset(request)

    def save_model(self, request, obj, form, change):
        super().save_model(request, obj, form, change)

@admin.register(Calendar)
class CalendarAdmin(admin.ModelAdmin):
    list_display = ('transport', 'title', 'start', 'end', 'status')
    list_filter = ('transport', 'status')
    search_fields = ('title', 'transport__name', 'transport__model')
    date_hierarchy = 'start'

    @staticmethod
    def calendar_view(request):
        from .models import Transport, Calendar
        transports = Transport.objects.all()
        events = Calendar.objects.all()
        if not request.user.is_superuser:
            profile = getattr(request.user, 'profile', None)
            if profile:
                transports = transports.filter(city=profile.city)
                events = events.filter(city=profile.city)
            else:
                transports = Transport.objects.none()
                events = Calendar.objects.none()
        # Получаем статистику по статусам заявок
        from django.db.models import Count
        status_stats = events.values('status').annotate(
            count=Count('id')
        ).order_by('status')
        
        # Формируем статистику для отображения
        stats = {}
        total_rentals = 0
        for stat in status_stats:
            status_display = dict(RentalApplication.STATUS_CHOICES).get(stat['status'], stat['status'])
            stats[stat['status']] = {
                'count': stat['count'],
                'display': status_display
            }
            total_rentals += stat['count']
        
        context = {
            'transports': transports,
            'title': 'Календарь аренды',
            'opts': Calendar._meta,
            'stats': stats,
            'total_rentals': total_rentals,
        }
        return render(request, 'admin/calendar.html', context)

    @staticmethod
    def calendar_events(request):
        try:
            transport_id = request.GET.get('transport_id')
            transport_ids = request.GET.getlist('transport_ids[]')  # Для множественного выбора
            start = request.GET.get('start')
            end = request.GET.get('end')

            # Определяем фильтр для транспорта
            if transport_ids:
                # Если переданы ID нескольких транспортов
                events = Calendar.objects.filter(transport_id__in=transport_ids)
            elif not transport_id or transport_id == 'all':
                # Если transport_id не указан или равен "all", показываем все события
                events = Calendar.objects.all()
            else:
                # Фильтруем по одному транспорту
                events = Calendar.objects.filter(transport_id=transport_id)

            # Фильтрация по городу менеджера
            if not request.user.is_superuser:
                profile = getattr(request.user, 'profile', None)
                if profile:
                    events = events.filter(city=profile.city)
                else:
                    events = events.none()

            if start:
                try:
                    # FullCalendar отправляет даты в формате YYYY-MM-DD
                    start_date = datetime.strptime(start.split('T')[0], '%Y-%m-%d')
                    events = events.filter(start__date__gte=start_date.date())
                except (ValueError, IndexError) as e:
                    return JsonResponse({'error': f'Invalid start date format: {str(e)}'}, status=400)
            
            if end:
                try:
                    # FullCalendar отправляет даты в формате YYYY-MM-DD
                    end_date = datetime.strptime(end.split('T')[0], '%Y-%m-%d')
                    events = events.filter(end__date__lte=end_date.date())
                except (ValueError, IndexError) as e:
                    return JsonResponse({'error': f'Invalid end date format: {str(e)}'}, status=400)

            events_data = []
            for event in events:
                color = {
                    RentalApplication.STATUS_RESERVED: '#ffc107',  # желтый
                    RentalApplication.STATUS_ACTIVE: '#28a745',    # зеленый
                    RentalApplication.STATUS_COMPLETED: '#17a2b8', # голубой
                    RentalApplication.STATUS_CANCELLED: '#dc3545', # красный
                }.get(event.status, '#6c757d')  # серый по умолчанию

                # Добавляем информацию о транспорте в заголовок события
                title = f"{event.title} (№{event.transport.number} - {event.transport.name} {event.transport.model})"

                # Получаем дополнительную информацию о заявке для tooltip
                extended_props = {
                    'transport': f"№{event.transport.number} - {event.transport.name} {event.transport.model}",
                    'transportNumber': event.transport.number,
                    'transportId': event.transport.id,
                    'status': event.status,
                    'statusDisplay': dict(RentalApplication.STATUS_CHOICES).get(event.status, event.status),
                }

                # Добавляем информацию о стоимости и депозите, если есть связанная заявка
                if event.rental_application:
                    rental = event.rental_application
                    extended_props.update({
                        'clientName': rental.full_name,
                        'clientPhone': str(rental.phone_number),
                        'cost': f"{rental.calculate_total_cost():,}",
                        'deposit': f"{int(rental.security_deposit or 0):,}",
                        'discount': f"{rental.discount}%",
                        'dailyRate': f"{rental.get_daily_rate():,}",
                        'rentalDays': rental.get_rental_days(),
                    })

                events_data.append({
                    'id': event.rental_application.id if event.rental_application else event.id,
                    'title': title,
                    'start': event.start.isoformat(),
                    'end': event.end.isoformat(),
                    'allDay': event.all_day,
                    'color': color,
                    'url': f'/admin/rentals/rentalapplication/{event.rental_application.id}/change/' if event.rental_application else None,
                    'extendedProps': extended_props,
                })

            return JsonResponse(events_data, safe=False)
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)

    def changelist_view(self, request, extra_context=None):
        extra_context = extra_context or {}
        extra_context['show_calendar_link'] = True
        return super().changelist_view(request, extra_context=extra_context)

    def get_queryset(self, request):
        qs = super().get_queryset(request)
        if request.user.is_superuser:
            return qs
        profile = getattr(request.user, 'profile', None)
        if profile:
            return qs.filter(city=profile.city)  # исправлено: было transport__city
        return qs.none()

    def save_model(self, request, obj, form, change):
        super().save_model(request, obj, form, change)

@staff_member_required
def return_calendar_view(request):
    from .models import Calendar, Transport
    transports = Transport.objects.all()
    events = Calendar.objects.all()
    if not request.user.is_superuser:
        profile = getattr(request.user, 'profile', None)
        if profile:
            transports = transports.filter(city=profile.city)
            events = events.filter(city=profile.city)
        else:
            transports = Transport.objects.none()
            events = Calendar.objects.none()
    context = {
        'transports': transports,
        'title': 'Календарь сдачи транспорта',
        'opts': Calendar._meta,
    }
    return render(request, 'admin/calendar_returns.html', context)

@staff_member_required
def return_calendar_events(request):
    from .models import Calendar, RentalApplication
    try:
        transport_id = request.GET.get('transport_id')
        start = request.GET.get('start')
        end = request.GET.get('end')

        events = Calendar.objects.all()
        if transport_id and transport_id != 'all':
            events = events.filter(transport_id=transport_id)

        # Фильтрация по городу менеджера
        if not request.user.is_superuser:
            profile = getattr(request.user, 'profile', None)
            if profile:
                events = events.filter(city=profile.city)
            else:
                events = events.none()

        if start:
            from datetime import datetime
            start_date = datetime.strptime(start.split('T')[0], '%Y-%m-%d')
            events = events.filter(end__date__gte=start_date.date())
        if end:
            from datetime import datetime
            end_date = datetime.strptime(end.split('T')[0], '%Y-%m-%d')
            events = events.filter(end__date__lte=end_date.date())

        events_data = []
        for event in events:
            color = {
                RentalApplication.STATUS_RESERVED: '#ffc107',
                RentalApplication.STATUS_ACTIVE: '#28a745',
                RentalApplication.STATUS_COMPLETED: '#17a2b8',
                RentalApplication.STATUS_CANCELLED: '#dc3545',
            }.get(event.status, '#6c757d')
            title = f"Возврат: №{event.transport.number} - {event.transport.name} {event.transport.model} ({event.title})"
            # Используем дату окончания аренды из заявки, если есть
            if event.rental_application and event.rental_application.rental_end_date:
                return_date = event.rental_application.rental_end_date
            else:
                return_date = event.end.date()
            events_data.append({
                'id': event.rental_application.id if event.rental_application else event.id,
                'title': title,
                'start': return_date.isoformat(),
                'end': return_date.isoformat(),
                'allDay': True,
                'color': color,
                'url': f'/admin/rentals/rentalapplication/{event.rental_application.id}/change/' if event.rental_application else None,
                'extendedProps': {
                    'transport': f"№{event.transport.number} - {event.transport.name} {event.transport.model}",
                    'transportNumber': event.transport.number,
                    'status': event.status,
                }
            })
        return JsonResponse(events_data, safe=False)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

@admin.register(Advantages)
class AdvantagesAdmin(SummernoteModelAdmin):
    summernote_fields = ('text',)
    list_display = ('number', 'name', 'slug', 'created_at')
    list_display_links = ('name',)
    list_editable = ('number',)
    list_filter = ('created_at',)
    search_fields = ('name', 'text')
    prepopulated_fields = {'slug': ('name',)}
    ordering = ('number',)
    
    fieldsets = (
        ('Основная информация', {
            'fields': ('name', 'slug', 'number', 'text')
        }),
        ('Медиа', {
            'fields': ('image',)
        }),
        ('Метаданные', {
            'fields': ('created_at', 'updated_at'),
            'classes': ('collapse',)
        }),
    )
    
    readonly_fields = ('created_at', 'updated_at')

@admin.register(Blog)
class BlogAdmin(SummernoteModelAdmin):
    summernote_fields = ('text',)
    list_display = ('title', 'category', 'created_at')
    search_fields = ('title', 'category', 'text')
    list_filter = ('category', 'created_at')
    prepopulated_fields = {'slug': ('title',)}
    readonly_fields = ('created_at', 'updated_at')

    fieldsets = (
        ('Основное', {
            'fields': ('title', 'slug', 'category')
        }),
        ('Медиа', {
            'fields': ('main_image', 'extra_image')
        }),
        ('Текст', {
            'fields': ('text',)
        }),
        ('Метаданные', {
            'fields': ('created_at', 'updated_at'),
            'classes': ('collapse',)
        }),
    )

@admin.register(Review)
class ReviewAdmin(admin.ModelAdmin):
    list_display = ('name', 'created_at')
    search_fields = ('name', 'text')
    prepopulated_fields = {'slug': ('name',)}
    readonly_fields = ('created_at', 'updated_at')
    fieldsets = (
        ('Основное', {
            'fields': ('name', 'slug')
        }),
        ('Контент', {
            'fields': ('text', 'image')
        }),
        ('Метаданные', {
            'fields': ('created_at', 'updated_at'),
            'classes': ('collapse',)
        }),
    )